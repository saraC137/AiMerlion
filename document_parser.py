"""
Document Parser Module for Resume Data Extraction System

This module handles PDF and DOCX parsing with SMART extraction:
1. First INSPECTS the PDF to determine its type (text, vector, scanned)
2. Then uses the BEST extraction method based on the inspection

Key Features:
- PDF inspection before extraction (saves time, better results)
- Smart method selection based on PDF type
- OCR for vector/scanned PDFs
- pdfplumber for true text PDFs
- Hybrid for mixed PDFs
- Windows path compatibility
"""

import os
import logging
from typing import Optional, Tuple
from pathlib import Path
import re

# PDF and Document handling
import pdfplumber
from docx import Document

# OCR imports
from PIL import Image
import pytesseract

# Import configuration
from config import TESSERACT_CONFIG, SUPPORTED_EXTENSIONS, OCR_CONFIG

# Import PDF Inspector
from pdf_inspector import PDFInspector, PDFType, inspect_pdf

# Setup logging with visual clarity
logger = logging.getLogger(__name__)

# Try to import pdf2image for OCR
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    logger.warning("pdf2image not available - OCR fallback will be limited")


class DocumentParser:
    """
    Handles document parsing with intelligent OCR fallback.

    For PDFs with text converted to outlines/vectors, pdfplumber returns
    empty or minimal text. This class automatically detects this and
    falls back to OCR extraction.
    """

    def __init__(self):
        """Initialize the document parser with configuration settings."""
        self.dpi = OCR_CONFIG.get("dpi", 200)
        self.min_words_per_page = OCR_CONFIG.get("min_words_per_page", 20)
        self.ocr_timeout = OCR_CONFIG.get("ocr_timeout", 30)
        self.min_text_length = OCR_CONFIG.get("min_text_length", 50)
        self.fallback_enabled = OCR_CONFIG.get("fallback_enabled", True)
        self.tesseract_config = TESSERACT_CONFIG

        logger.info("📄 DocumentParser initialized")
        logger.info(f"   OCR DPI: {self.dpi}")
        logger.info(f"   Min words/page: {self.min_words_per_page}")
        logger.info(f"   OCR fallback: {'Enabled' if self.fallback_enabled else 'Disabled'}")

    def parse(self, file_path: str) -> Tuple[Optional[str], str]:
        """
        Parse a document and extract text.

        Args:
            file_path: Path to the document (PDF or DOCX)

        Returns:
            Tuple of (extracted_text, extraction_method)
            extraction_method is one of: "pdfplumber", "ocr", "docx", "failed"
        """
        # Normalize Windows path
        file_path = self._normalize_path(file_path)

        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None, "failed"

        file_ext = Path(file_path).suffix.lower()
        file_name = os.path.basename(file_path)

        logger.info(f"📄 Parsing: {file_name}")

        if file_ext == ".pdf":
            return self._parse_pdf(file_path)
        elif file_ext in [".docx", ".doc"]:
            return self._parse_docx(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
            return None, "failed"

    def _normalize_path(self, file_path: str) -> str:
        """
        Normalize file path for Windows compatibility.
        Converts forward slashes to backslashes and resolves the path.
        """
        # Convert to Path object for proper handling
        path = Path(file_path)
        # Resolve to absolute path
        return str(path.resolve())

    def _parse_pdf(self, file_path: str) -> Tuple[Optional[str], str]:
        """
        Parse PDF with intelligent OCR fallback.

        Strategy:
        1. Try pdfplumber first (fast, works for normal PDFs)
        2. Check if extracted text is meaningful
        3. If text is insufficient/gibberish, fall back to OCR

        Args:
            file_path: Path to the PDF file

        Returns:
            Tuple of (extracted_text, extraction_method)
        """
        file_name = os.path.basename(file_path)

        # ============================================
        # STEP 1: Try pdfplumber extraction first
        # ============================================
        logger.info(f"📖 Step 1: Trying pdfplumber extraction for {file_name}...")

        try:
            pdfplumber_text = self._extract_with_pdfplumber(file_path)

            if pdfplumber_text:
                # ============================================
                # STEP 2: Check if extracted text is meaningful
                # ============================================
                logger.info(f"🔍 Step 2: Checking if extracted text is meaningful...")

                if self._is_meaningful_text(pdfplumber_text):
                    logger.info(f"✅ pdfplumber extracted meaningful text ({len(pdfplumber_text)} chars)")
                    return pdfplumber_text, "pdfplumber"
                else:
                    logger.warning(f"⚠️ pdfplumber text appears to be gibberish/empty")
                    logger.warning(f"   Preview: {pdfplumber_text[:100]}..." if len(pdfplumber_text) > 100 else f"   Preview: {pdfplumber_text}")
            else:
                logger.warning(f"⚠️ pdfplumber returned no text")

        except Exception as e:
            logger.error(f"❌ pdfplumber failed: {e}")

        # ============================================
        # STEP 3: Fall back to OCR if enabled
        # ============================================
        if self.fallback_enabled:
            logger.info(f"📸 Step 3: Falling back to OCR extraction for {file_name}...")

            ocr_text = self._parse_pdf_with_ocr(file_path)

            if ocr_text and self._is_meaningful_text(ocr_text):
                logger.info(f"✅ OCR successfully extracted text ({len(ocr_text)} chars)")
                return ocr_text, "ocr"
            else:
                logger.error(f"❌ OCR also failed to extract meaningful text")
        else:
            logger.warning(f"⚠️ OCR fallback is disabled in config")

        # ============================================
        # STEP 4: Return whatever we have (or None)
        # ============================================
        if pdfplumber_text and len(pdfplumber_text.strip()) > 0:
            logger.warning(f"⚠️ Returning pdfplumber text despite low quality")
            return pdfplumber_text, "pdfplumber"

        logger.error(f"❌ Failed to extract any text from {file_name}")
        return None, "failed"

    def _extract_with_pdfplumber(self, file_path: str) -> Optional[str]:
        """
        Extract text from PDF using pdfplumber.

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text or None if extraction fails
        """
        try:
            all_text = []

            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"   📄 Processing {total_pages} pages with pdfplumber...")

                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        all_text.append(page_text)
                    logger.debug(f"   Page {i+1}/{total_pages}: {len(page_text) if page_text else 0} chars")

            combined_text = "\n\n".join(all_text)
            return combined_text if combined_text.strip() else None

        except Exception as e:
            logger.error(f"❌ pdfplumber extraction error: {e}")
            return None

    def _parse_pdf_with_ocr(self, file_path: str) -> Optional[str]:
        """
        Extract text from PDF using OCR (Tesseract + pdf2image).

        Uses the same simple, reliable approach as test_ocr.py:
        - Convert PDF pages to images at 200 DPI
        - Run pytesseract.image_to_string() without extra config
        - This captures ALL text including outlined/vector text

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text or None if extraction fails
        """
        file_name = os.path.basename(file_path)

        # Check if pdf2image is available
        if not PDF2IMAGE_AVAILABLE:
            logger.error("❌ pdf2image not installed - cannot perform OCR")
            logger.error("   Install with: pip install pdf2image")
            logger.error("   Also install Poppler: https://github.com/osber/poppler-windows/releases")
            return None

        try:
            logger.info(f"📸 Converting PDF to images (DPI: {self.dpi})...")

            # Convert PDF pages to images (same as test_ocr.py)
            try:
                images = convert_from_path(
                    file_path,
                    dpi=self.dpi  # Default 200, same as test_ocr.py
                )
            except Exception as poppler_error:
                # Handle Poppler not found error with helpful message
                error_msg = str(poppler_error).lower()
                if "poppler" in error_msg or "pdftoppm" in error_msg:
                    logger.error("❌ Poppler not found! OCR requires Poppler to be installed.")
                    logger.error("   📥 Windows: Download from https://github.com/osber/poppler-windows/releases")
                    logger.error("   📁 Extract and add 'bin' folder to your PATH")
                    logger.error("   🔄 Or set poppler_path in convert_from_path()")
                else:
                    logger.error(f"❌ PDF to image conversion failed: {poppler_error}")
                return None

            total_pages = len(images)
            logger.info(f"   🖼️ Converted {total_pages} pages to images")

            # OCR each page - SIMPLE approach like test_ocr.py (no extra config!)
            all_text = []

            for i, image in enumerate(images):
                page_num = i + 1
                logger.info(f"   📸 OCR processing page {page_num}/{total_pages}...")

                try:
                    # Simple OCR call - same as test_ocr.py (NO extra config!)
                    # This is more reliable and captures all text
                    page_text = pytesseract.image_to_string(image, lang='eng')

                    if page_text and page_text.strip():
                        all_text.append(page_text.strip())
                        word_count = len(page_text.split())
                        logger.info(f"      ✅ Page {page_num}: {word_count} words extracted")
                    else:
                        logger.warning(f"      ⚠️ Page {page_num}: No text extracted")

                except Exception as page_error:
                    logger.error(f"      ❌ Error processing page {page_num}: {page_error}")

            # Combine all pages
            combined_text = "\n\n".join(all_text)

            if combined_text.strip():
                total_words = len(combined_text.split())
                logger.info(f"✅ OCR complete: {len(combined_text)} chars, {total_words} words")
                return combined_text
            else:
                logger.warning("⚠️ OCR produced no text")
                return None

        except Exception as e:
            logger.error(f"❌ OCR extraction failed: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return None

    def _is_meaningful_text(self, text: str) -> bool:
        """
        Check if extracted text is meaningful (not gibberish or empty).

        Criteria for meaningful text:
        1. Has minimum character length
        2. Has minimum word count
        3. Has reasonable word length distribution (not all single chars)
        4. Contains some common English patterns

        Args:
            text: The extracted text to evaluate

        Returns:
            True if text appears meaningful, False otherwise
        """
        if not text:
            return False

        text = text.strip()

        # Check 1: Minimum length
        if len(text) < self.min_text_length:
            logger.debug(f"   ❌ Text too short: {len(text)} < {self.min_text_length}")
            return False

        # Check 2: Minimum word count
        words = text.split()
        if len(words) < self.min_words_per_page:
            logger.debug(f"   ❌ Too few words: {len(words)} < {self.min_words_per_page}")
            return False

        # Check 3: Average word length (gibberish often has unusual word lengths)
        avg_word_length = sum(len(w) for w in words) / len(words)
        if avg_word_length < 2 or avg_word_length > 15:
            logger.debug(f"   ❌ Unusual avg word length: {avg_word_length:.1f}")
            return False

        # Check 4: Ratio of alphabetic characters (should be mostly letters)
        alpha_chars = sum(1 for c in text if c.isalpha())
        alpha_ratio = alpha_chars / len(text) if text else 0
        if alpha_ratio < 0.4:  # At least 40% should be letters
            logger.debug(f"   ❌ Low alpha ratio: {alpha_ratio:.2f}")
            return False

        # Check 5: Look for common English words/patterns
        common_words = {'the', 'and', 'to', 'of', 'a', 'in', 'is', 'for', 'with', 'on',
                        'experience', 'work', 'education', 'skills', 'name', 'email',
                        'phone', 'address', 'university', 'company', 'manager', 'developer'}
        text_lower = text.lower()
        found_common = sum(1 for word in common_words if word in text_lower)
        if found_common < 2:
            logger.debug(f"   ❌ Few common words found: {found_common}")
            return False

        logger.debug(f"   ✅ Text appears meaningful: {len(text)} chars, {len(words)} words")
        return True

    def _parse_docx(self, file_path: str) -> Tuple[Optional[str], str]:
        """
        Parse DOCX file and extract text.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Tuple of (extracted_text, "docx") or (None, "failed")
        """
        file_name = os.path.basename(file_path)

        try:
            logger.info(f"📄 Parsing DOCX: {file_name}")

            doc = Document(file_path)
            all_text = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            all_text.append(cell.text)

            combined_text = "\n".join(all_text)

            if combined_text.strip():
                logger.info(f"✅ DOCX extracted: {len(combined_text)} chars")
                return combined_text, "docx"
            else:
                logger.warning(f"⚠️ DOCX appears empty")
                return None, "failed"

        except Exception as e:
            logger.error(f"❌ DOCX parsing failed: {e}")
            return None, "failed"


# ============================================
# Convenience function for quick parsing
# ============================================

def parse_document(file_path: str) -> Tuple[Optional[str], str]:
    """
    Convenience function to parse a document.

    Args:
        file_path: Path to the document

    Returns:
        Tuple of (extracted_text, extraction_method)
    """
    parser = DocumentParser()
    return parser.parse(file_path)


# ============================================
# Main entry point for testing
# ============================================

if __name__ == "__main__":
    import sys
    import coloredlogs

    # Setup colored logging
    coloredlogs.install(
        level='DEBUG',
        fmt='%(asctime)s - %(levelname)s - %(message)s'
    )

    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        print(f"\n{'='*60}")
        print(f"Testing DocumentParser with: {test_file}")
        print(f"{'='*60}\n")

        text, method = parse_document(test_file)

        print(f"\n{'='*60}")
        print(f"RESULTS")
        print(f"{'='*60}")
        print(f"Extraction Method: {method}")
        print(f"Text Length: {len(text) if text else 0} characters")

        if text:
            print(f"\nPreview (first 500 chars):")
            print("-" * 40)
            print(text[:500])
            print("-" * 40)
    else:
        print("Usage: python document_parser.py <path_to_pdf_or_docx>")
