# 🧙‍♀️✨ Fairy Codemother's ULTIMATE Resume Extractor - THE ID DIVA EDITION! ✨🧙‍♀️

import numpy as np
import os
import re
import datetime
import pandas as pd
from tqdm import tqdm
import logging
import coloredlogs
import ollama
import pypdf
import docx
import pdfplumber
from pathlib import Path
from paddleocr import PaddleOCR
from PIL import Image
import pdf2image
from typing import Dict, List, Tuple, Optional
import math
import time
import unicodedata
import json
import config
from utils import display_menu, save_checkpoint, load_checkpoint, print_batch_table, FeedbackLoopSystem, InteractiveCorrectionSystem, PatternLearningSystem, PerformanceMonitor, standardize_phone_number, standardize_date
from ai_validator import AIValidator

# Setup logging
logger = logging.getLogger(__name__)
coloredlogs.install(level='INFO', logger=logger, fmt='%(asctime)s - 💄 %(levelname)s - %(message)s')

def find_candidate_folders(base_folder: str) -> List[str]:
    """
    🔍 Find all candidate folders within the base directory structure.
    Handles structure like: merlion_resumes/2025-07-10_16-54-36/[candidate_folders]
    """
    candidate_folders = []
    
    # First, check if base_folder has date-like subfolders
    for item in os.scandir(base_folder):
        if item.is_dir() and not item.name.startswith('.'):
            # Check if this is a date folder (like 2025-07-10_16-54-36)
            if re.match(r'\d{4}-\d{2}-\d{2}', item.name):
                logger.info(f"📁 Found date folder: {item.name}")
                # Look for candidate folders inside
                for candidate in os.scandir(item.path):
                    if candidate.is_dir() and not candidate.name.startswith('.'):
                        # Any folder inside the date folder is a candidate folder
                        logger.debug(f"   📂 Found candidate folder: {candidate.name}")
                        candidate_folders.append(candidate.path)
            # Or if it's directly a candidate folder (fallback)
            elif re.match(r'^\d+', item.name):
                candidate_folders.append(item.path)
    
    logger.info(f"🎯 Total candidate folders found: {len(candidate_folders)}")
    return sorted(candidate_folders)


class UltimateResumeExtractor:
    """
    🎭 THE ULTIMATE DIVA! This queen relies on POWERFUL regex patterns
    and strategic AI usage. She's learned that sometimes, you need to
    do the heavy lifting yourself instead of relying on others!
    """


    def __init__(self, model_name: str):
        self.model_name = model_name
        
        # ✨ REMOVED: Japanese OCR (PaddleOCR)
        # self.ocr = PaddleOCR(use_angle_cls=True, lang='japan')
        
        # 🆕 ADD: English OCR (much faster!)
        try:
            import pytesseract
            self.ocr_available = True
            logger.info("✨ English OCR ready!")
        except ImportError:
            self.ocr_available = False
            logger.warning("⚠️ pytesseract not available - OCR disabled")
        
        # AI setup
        self.use_ai = config.USE_AI_EXTRACTION
        self.ai_extractor = None
        self.ai_enabled = False
        
        if self.use_ai:
            try:
                from ai_extractor import AIExtractor
                self.ai_extractor = AIExtractor(model_name, logger=logger)
                self.ai_enabled = self.ai_extractor.available
                if self.ai_enabled:
                    logger.info("🌸 AI extraction ENABLED!")
                else:
                    logger.warning("⚠️ AI not available - using regex only")
            except Exception as e:
                logger.error(f"❌ Could not initialize AI: {e}")
        else:
            logger.info("🎯 AI extraction disabled - using regex only")

    def process_individual_file(self, file_path: str) -> Dict:
        """
        🎯 Process a single resume file as an individual candidate
        """
        logger.info(f"📄 Processing individual file: {os.path.basename(file_path)}")
        
        # Extract ID from filename if possible
        filename = os.path.basename(file_path)
        id_match = re.search(r'(\d+)', filename)
        
        result = {
            "ID": int(id_match.group(1)) if id_match else None,
            "Name": None, 
            "Email": None, 
            "Phone": None,
            "Date_of_Birth": None, 
            "Extraction_Status": "Failed",
            "Notes": "",
            "AI_Assisted": False,
            "Filenames_Processed": filename
        }
        
        # Get text and extract data
        text = self.get_text_from_file(file_path)
        if text and len(text.strip()) >= 50:
            extracted_data, ai_used = self._extract_data_from_text(text)
            
            if ai_used:
                result["AI_Assisted"] = True
            
            # Update result with extracted data
            for field in ["Name", "Email", "Phone", "Date_of_Birth"]:
                if extracted_data.get(field):
                    result[field] = extracted_data[field]
            
            # Try to extract name from filename if still missing
            if not result["Name"]:
                name_from_file = self._extract_name_from_filename(filename)
                if name_from_file:
                    result["Name"] = name_from_file
                    result["Notes"] += "Name extracted from filename. "
        
        self._set_extraction_status(result, text if text else "")
        return result

    def _apply_learned_patterns(self):
        """
        🧠 Apply patterns we've learned from corrections!
        """
        patterns = self.feedback_system.load_learned_patterns()
        
        # Check if we have significant learnings
        total_corrections = sum(
            len(patterns.get(field, {}).get('transformations', []))
            for field in patterns
        )
        
        if total_corrections > 10:
            logger.info(f"🧠 Loaded {total_corrections} learned patterns!")
            
            # Suggest improvements
            suggestions = self.pattern_learner.generate_regex_suggestions()
            for field, field_suggestions in suggestions.items():
                if field_suggestions:
                    logger.info(f"💡 Suggestions for {field}: {field_suggestions[0]}")

    def _extract_with_mega_regex(self, text: str) -> Dict[str, Optional[str]]:
        """
        🌟 MEGA REGEX for ENGLISH resumes only!
        Cleaner, faster, more accurate!
        """
        logger.debug("🎭 ENGLISH-ONLY REGEX BEGINS!")
        
        data: Dict[str, Optional[str]] = {
            "name": None,
            "email": None,
            "phone": None,
            "date_of_birth": None,
            "skills": None,
            "working_experience": None,
            "location": None,
            "school_university": None
        }
        
        # Try table extraction first
        table_data = self._extract_from_table_format(text)
        data.update(table_data)
        
        # --- 📧 EMAIL EXTRACTION (most reliable!) ---
        if not data["email"]:
            email_patterns = [
                r'\b[A-Za-z0-9][A-Za-z0-9._%+-]*@[A-Za-z0-9][A-Za-z0-9.-]*\.[A-Z|a-z]{2,}\b',
                r'(?:[Ee]-?[Mm]ail|MAIL|Email)[\s:]*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',
            ]
            for pattern in email_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    email = matches[0] if isinstance(matches[0], str) else matches[0]
                    data["email"] = email.lower()
                    logger.info(f"✨ Found email: {data['email']}")
                    break
        
        # --- 📱 PHONE EXTRACTION ---
        if not data["phone"]:
            phone = self._extract_phone_english(text, data.get("email"))
            if phone:
                data["phone"] = phone
        
        # --- 🎂 DATE OF BIRTH ---
        if not data["date_of_birth"]:
            dob = self._extract_dob_english(text)
            if dob:
                data["date_of_birth"] = dob
        
        # --- 👤 NAME EXTRACTION ---
        if not data["name"]:
            name = self._extract_name_english(text)
            if name:
                data["name"] = name
        
        # --- 🎯 SKILLS ---
        if not data["skills"]:
            skills_match = re.search(r"(?:Skills|SKILLS|Technical Skills|Core Competencies)[\s:]*([^\n\r]+(?:\n(?![\n])[^\n]+)*)", 
                                    text, re.IGNORECASE)
            if skills_match:
                data["skills"] = skills_match.group(1).strip()
        
        # --- 💼 WORK EXPERIENCE ---
        if not data["working_experience"]:
            exp_match = re.search(r"(?:Experience|EXPERIENCE|Work Experience|Employment History)[\s:]*([^\n\r]+(?:\n(?![\n])[^\n]+)*)", 
                                text, re.IGNORECASE)
            if exp_match:
                data["working_experience"] = exp_match.group(1).strip()
        
        # --- 📍 LOCATION ---
        if not data["location"]:
            location_match = re.search(r"(?:Location|Address|City|Residence)[\s:]*([^\n\r]+)", 
                                    text, re.IGNORECASE)
            if location_match:
                data["location"] = location_match.group(1).strip()
        
        # --- 🎓 EDUCATION ---
        if not data["school_university"]:
            edu_match = re.search(r"(?:Education|EDUCATION|University|College|School)[\s:]*([^\n\r]+)", 
                                text, re.IGNORECASE)
            if edu_match:
                data["school_university"] = edu_match.group(1).strip()
        
        return data

    def _extract_phone_english(self, text: str, email: Optional[str]) -> Optional[str]:
        """
        📱 Extract phone numbers from ENGLISH resumes!
        Supports US, UK, and international formats.
        """
        
        # If we have an email, search near it first
        search_areas = []
        if email:
            email_index = text.find(email)
            if email_index != -1:
                start = max(0, email_index - 300)
                end = min(len(text), email_index + 300)
                search_areas.append(("near email", text[start:end]))
        
        search_areas.append(("full text", text))
        
        # Phone patterns for English resumes
        phone_patterns = [
            # US/Canada: (123) 456-7890
            (r'\((\d{3})\)\s*(\d{3})[-\s]?(\d{4})', 'us_paren'),
            # US/Canada: 123-456-7890
            (r'\b(\d{3})[-\.\s](\d{3})[-\.\s](\d{4})\b', 'us_dash'),
            # International: +1 123 456 7890
            (r'\+1\s*\(?(\d{3})\)?[-\.\s]?(\d{3})[-\.\s]?(\d{4})', 'us_intl'),
            # UK: +44 1234 567890
            (r'\+44\s*(\d{3,4})\s*(\d{3})\s*(\d{3,4})', 'uk'),
            # After label
            (r'(?:Phone|Tel|Mobile|Cell|Contact)[\s:]*([+\d\-\s\(\)\.]{10,20})', 'labeled'),
            # Generic 10+ digits
            (r'\b(\d{3}[-\.\s]?\d{3}[-\.\s]?\d{4})\b', 'generic'),
        ]
        
        for area_name, search_text in search_areas:
            for pattern, pattern_type in phone_patterns:
                matches = re.finditer(pattern, search_text, re.IGNORECASE)
                for match in matches:
                    phone_text = match.group(0) if pattern_type == 'labeled' else match.group(0)
                    
                    # Validate it looks like a phone
                    digits_only = re.sub(r'\D', '', phone_text)
                    
                    # Must be 10-15 digits
                    if 10 <= len(digits_only) <= 15:
                        # Not all same digit
                        if not re.match(r'^(\d)\1+$', digits_only):
                            formatted = standardize_phone_number(phone_text)
                            if formatted:
                                logger.info(f"✨ Found phone: {formatted}")
                                return formatted
        
        return None

    def _extract_dob_english(self, text: str) -> Optional[str]:
        """
        🎂 Extract date of birth from ENGLISH resumes!
        Returns YYYY-MM-DD format.
        """
        import datetime
        
        current_year = datetime.datetime.now().year
        min_birth_year = current_year - config.MAX_AGE
        max_birth_year = current_year - config.MIN_AGE
        
        # Find contact info area
        contact_area = self._find_contact_area(text)
        search_text = contact_area if contact_area else text
        
        # DOB patterns for English resumes
        dob_patterns = [
            # With label: DOB: 01/15/1990
            (r'(?:DOB|D\.O\.B\.|Date of Birth|Birth Date|Born)[\s:]*(\d{1,2})[/-](\d{1,2})[/-](\d{4})', 'mdy_labeled'),
            # With label: DOB: 1990-01-15
            (r'(?:DOB|D\.O\.B\.|Date of Birth|Birth Date|Born)[\s:]*(\d{4})[/-](\d{1,2})[/-](\d{1,2})', 'ymd_labeled'),
            # Written: January 15, 1990
            (r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{1,2}),?\s+(\d{4})\b', 'written'),
            # ISO format: 1990-01-15
            (r'\b(\d{4})-(\d{1,2})-(\d{1,2})\b', 'ymd'),
            # US format: 01/15/1990
            (r'\b(\d{1,2})/(\d{1,2})/(\d{4})\b', 'mdy'),
        ]
        
        for pattern, date_format in dob_patterns:
            matches = re.finditer(pattern, search_text, re.IGNORECASE)
            for match in matches:
                try:
                    if date_format == 'written':
                        month_name, day, year = match.groups()
                        month = datetime.datetime.strptime(month_name, '%B').month
                        year, day = int(year), int(day)
                    
                    elif date_format in ['mdy_labeled', 'mdy']:
                        month, day, year = map(int, match.groups())
                    
                    elif date_format in ['ymd_labeled', 'ymd']:
                        year, month, day = map(int, match.groups())
                    
                    # Validate year range
                    if min_birth_year <= year <= max_birth_year:
                        # Validate date is real
                        dt = datetime.datetime(year, month, day)
                        formatted = dt.strftime('%Y-%m-%d')
                        logger.info(f"✅ Found DOB: {formatted}")
                        return formatted
                
                except (ValueError, OverflowError):
                    continue
        
        return None

    def _find_contact_area(self, text: str) -> Optional[str]:
        """
        🔍 Find the contact information section
        """
        # Look for email as anchor
        email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        if email_match:
            pos = email_match.start()
            start = max(0, pos - 500)
            end = min(len(text), pos + 500)
            return text[start:end]
        
        # Look for phone as anchor
        phone_match = re.search(r'\b\d{3}[-\.\s]\d{3}[-\.\s]\d{4}\b', text)
        if phone_match:
            pos = phone_match.start()
            start = max(0, pos - 500)
            end = min(len(text), pos + 500)
            return text[start:end]
        
        return None

    def _extract_phone(self, text: str, email: Optional[str]) -> Optional[str]:
        """
        📱 ENHANCED Phone extraction with VERTICAL number support!
        Now handles those DRAMATIC vertical layouts!
        """
        # First, let's fix potential vertical phone numbers
        text = self._fix_vertical_phone_numbers(text)
        
        email_vicinity_text = ""
        if email:
            # Get text around the email (±300 characters for better coverage)
            email_index = text.find(email)
            if email_index != -1:
                start = max(0, email_index - 300)
                end = min(len(text), email_index + 300)
                email_vicinity_text = text[start:end]
                logger.debug(f"🔍 Searching for phone near email at index {email_index}")

        # Comprehensive phone patterns - GLOBAL EDITION!
        phone_patterns = [
            # Japanese mobile phones (090, 080, 070, 050)
            (r'(?:0[5789]0)[\-\s\(\)\.・－]*(\d{4})[\-\s\(\)\.・－]*(\d{4})', 'JP'),
            # Japanese with parentheses like (090) 1234-5678
            (r'\(0[5789]0\)\s*(\d{4})[\-\s・－]*(\d{4})', 'JP'),
            # Full-width Japanese
            (r'[（(]?[０0][５5-９9][０0][）)]?[\s\-－・]*([０-９0-9]{4})[\s\-－・]*([０-９0-9]{4})', 'JP'),
            # After headers (any language)
            (r'(?:電話|携帯電話|TEL|Tel|Phone|Mobile|Cell|携帯)[\s:：]*([（(]?[+\d\-\s\(\)\.・－０-９]{10,20}[）)]?)', 'HEADER'),
            # International format with +81
            (r'(?:\+81|８１)[\-\s\(\)]*(\d{1,2})[\-\s\(\)]*(\d{4})[\-\s\(\)]*(\d{4})', 'JP_INTL'),
            # Generic number sequences
            (r'(?:^|\s)([+\d\-\s\(\)\.]{10,20})(?:\s|$)', 'GENERIC'),
        ]

        # Try email vicinity first if we have an email
        search_areas = []
        if email_vicinity_text:
            search_areas.append(("email vicinity", email_vicinity_text))
        search_areas.append(("full text", text))
        
        for area_name, search_text in search_areas:
            logger.debug(f"🎯 Searching in {area_name}")
            
            for pattern, pattern_type in phone_patterns:
                matches = re.finditer(pattern, search_text, re.IGNORECASE | re.MULTILINE)
                for match in matches:
                    phone = self._extract_and_validate_phone(match, pattern_type)
                    if phone:
                        # Standardize the format
                        phone = self._standardize_phone_format(phone, pattern_type)
                        logger.info(f"✨ Found phone in {area_name}: {phone} (type: {pattern_type})")
                        return phone
        
        # EMERGENCY: Look for any 10+ digit number
        logger.warning("🚨 No standard phone found - trying desperate measures!")
        desperate_pattern = r'(?:^|\s)(\d[\d\-\s\(\)\.]{9,19})(?:\s|$)'
        matches = re.findall(desperate_pattern, text)
        for match in matches:
            digits_only = re.sub(r'\D', '', match)
            if 10 <= len(digits_only) <= 15:
                if not re.match(r'^(\d)\1+$', digits_only):  # Not all same digit
                    phone = self._standardize_phone_format(match, 'EMERGENCY')
                    logger.info(f"✨ Found phone (desperate mode): {phone}")
                    return phone
        
        return None

    def _fix_vertical_phone_numbers(self, text: str) -> str:
        """
        🔧 Fix vertical phone numbers like:
        (090)
        6074-6688
        """
        # Pattern for vertical phone numbers
        vertical_patterns = [
            # (090)\n1234-5678
            r'\((\d{3})\)\s*\n\s*(\d{4}[-\s]\d{4})',
            # (090)\n1234\n5678
            r'\((\d{3})\)\s*\n\s*(\d{4})\s*\n\s*(\d{4})',
            # 090\n1234\n5678
            r'(\d{3})\s*\n\s*(\d{4})\s*\n\s*(\d{4})',
            # Full-width versions
            r'[（(]([０-９0-9]{3})[）)]\s*\n\s*([０-９0-9]{4}[-\s－][０-９0-9]{4})',
        ]
        
        for pattern in vertical_patterns:
            def replacer(match):
                if len(match.groups()) == 2:
                    return f"({match.group(1)}) {match.group(2)}"
                elif len(match.groups()) == 3:
                    return f"({match.group(1)}) {match.group(2)}-{match.group(3)}"
                return match.group(0)
            
            text = re.sub(pattern, replacer, text, flags=re.MULTILINE)
        
        return text

    def _standardize_phone_format(self, phone: str, pattern_type: str) -> str:
        """
        📱 Standardize phone to Japanese or International format
        Remove + from international format as requested!
        """
        # First normalize full-width to half-width
        phone = phone.translate(str.maketrans('０１２３４５６７８９', '0123456789'))
        phone = phone.translate(str.maketrans('（）', '()'))
        
        # Extract just digits
        digits = re.sub(r'\D', '', phone)
        
        if pattern_type in ['JP', 'JP_INTL'] or (pattern_type == 'GENERIC' and digits.startswith('0')):
            # Japanese format
            if digits.startswith('81'):  # Remove country code
                digits = '0' + digits[2:]
            
            if len(digits) == 11 and digits.startswith('0'):
                # Format as 090-1234-5678
                return f"{digits[:3]}-{digits[3:7]}-{digits[7:]}"
            elif len(digits) == 10 and digits.startswith('0'):
                # Landline format 03-1234-5678
                if digits[1] in '3456789':  # Major cities
                    return f"{digits[:2]}-{digits[2:6]}-{digits[6:]}"
                else:
                    return f"{digits[:3]}-{digits[3:6]}-{digits[6:]}"
        
        # For international numbers, remove the + as requested
        if digits.startswith('81'):
            return f"81-{digits[2:4]}-{digits[4:8]}-{digits[8:]}"
        
        # Default format for other numbers
        if len(digits) >= 10:
            # Just group nicely
            if len(digits) == 10:
                return f"{digits[:3]}-{digits[3:6]}-{digits[6:]}"
            elif len(digits) == 11:
                return f"{digits[:3]}-{digits[3:7]}-{digits[7:]}"
            else:
                return phone.strip()  # Return as is
        
        return phone.strip()

    def _extract_english_name(self, text: str) -> Optional[str]:
        """Extract English names only"""
        patterns = [
            r'(?:Name|Full\s*)?Name[\s:：]*([A-Za-z\s\'.-]{2,50})',
            r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}$',  # Standard English name
        ]
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                name = match.strip()
                if self._is_valid_name_strict(name):
                    return name
        return None

    def _extract_japanese_name(self, text: str) -> Optional[str]:
        """Extract Japanese names only"""
        patterns = [
            r'氏\s*名[\s:：]*([一-龯]{1,4}[\s　]*[一-龯]{1,4})',
            r'([一-龯]{1,4}[\s　]*[一-龯]{1,4})[\s　]*[（(]([ぁ-ゖァ-ヾ\s　]+)[）)]',
            r'^[一-龯]{2,4}[\s　]+[一-龯]{2,4}$',
        ]
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                # If the pattern has groups, the result is a tuple
                name = (match[0] if isinstance(match, tuple) else match).strip()
                if self._is_valid_name_strict(name):
                    return name
        return None

    def _normalize_phone(self, phone: str) -> str:
        phone = phone.translate(str.maketrans('０１２３４５６７８９', '0123456789'))
        phone = phone.replace('－', '-').replace('・', '-')
        return re.sub(r'\s+', '', phone).strip()

    def _extract_dob(self, text: str) -> Optional[str]:
        """
        🎂 DOB Detective - Now searches near other contact info!
        ALWAYS returns YYYY-MM-DD format!
        """
        logger.debug("🎂 DOB Detective is on the case!")
        
        current_year = datetime.datetime.now().year
        min_birth_year = current_year - 65  # Max age 65
        max_birth_year = current_year - 21  # Min age 21
        
        def normalize_numbers(text_str: str) -> str:
            return text_str.translate(str.maketrans('０１２３４５６７８９', '0123456789'))
        
        # First, try to find DOB near email/phone if we have them
        contact_area = self._find_contact_information_area(text)
        if contact_area:
            logger.debug("🎯 Found contact area, searching for DOB there first!")
            dob = self._extract_dob_from_text(contact_area, min_birth_year, max_birth_year)
            if dob:
                return dob
        
        # If not found in contact area, search whole text
        return self._extract_dob_from_text(text, min_birth_year, max_birth_year)

    def _extract_dob_from_text(self, text: str, min_year: int, max_year: int) -> Optional[str]:
        """Extract DOB from given text section"""
        
        def normalize_numbers(text_str: str) -> str:
            return text_str.translate(str.maketrans('０１２３４５６７８９', '0123456789'))
        
        dob_patterns = [
            # Western style dates
            (r'(?:Date of Birth|DOB|Born|Birthday)[\s:：]*(\d{4})[/\-\.](\d{1,2})[/\-\.](\d{1,2})', 'ymd'),
            (r'(?:Date of Birth|DOB|Born|Birthday)[\s:：]*(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})', 'mdy'),
            
            # Japanese dates with headers
            (r'生年月日[\s:：]*(?:大正|昭和|平成|令和)?\s*([０-９0-9]{1,2})\s*年\s*([０-９0-9]{1,2})\s*月\s*([０-９0-9]{1,2})\s*日', 'era'),
            (r'生年月日[\s:：]*([０-９0-9]{4})\s*年\s*([０-９0-9]{1,2})\s*月\s*([０-９0-9]{1,2})\s*日', 'ymd'),
            
            # Dates with 生 marker
            (r'([０-９0-9]{4})年([０-９0-9]{1,2})月([０-９0-9]{1,2})日生', 'ymd'),
            (r'((?:19|20)\d{2})年(\d{1,2})月(\d{1,2})日生', 'ymd'),
            
            # ISO format
            (r'\b((?:19|20)\d{2})-(\d{1,2})-(\d{1,2})\b', 'ymd'),
            
            # Just year/month/day with various separators
            (r'\b((?:19|20)\d{2})[/\-\.](\d{1,2})[/\-\.](\d{1,2})\b', 'ymd'),
        ]
        
        for pattern, date_format in dob_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                try:
                    year, month, day = None, None, None
                    
                    if date_format == 'era' and len(match) == 3:
                        # Handle Japanese era
                        era_year = int(normalize_numbers(match[0]))
                        month = int(normalize_numbers(match[1]))
                        day = int(normalize_numbers(match[2]))
                        
                        # Detect era from context
                        context = text[max(0, text.find(match[0])-50):text.find(match[0])+50]
                        if '令和' in context:
                            year = 2018 + era_year
                        elif '平成' in context:
                            year = 1988 + era_year
                        elif '昭和' in context:
                            year = 1925 + era_year
                        else:
                            # Default to Heisei for reasonable ages
                            year = 1988 + era_year
                    
                    elif date_format == 'ymd':
                        year = int(normalize_numbers(match[0]))
                        month = int(normalize_numbers(match[1]))
                        day = int(normalize_numbers(match[2]))
                    
                    elif date_format == 'mdy':
                        month = int(normalize_numbers(match[0]))
                        day = int(normalize_numbers(match[1]))
                        year = int(normalize_numbers(match[2]))
                    
                    # Validate year
                    if year and min_year <= year <= max_year:
                        # Always format as YYYY-MM-DD
                        formatted_date = f"{year:04d}-{month:02d}-{day:02d}"
                        logger.info(f"✅ Found valid DOB: {formatted_date}")
                        return formatted_date
                        
                except (ValueError, IndexError) as e:
                    logger.debug(f"Could not parse date: {match} - {e}")
                    continue
        
        return None

    def _find_contact_information_area(self, text: str) -> Optional[str]:
        """
        🔍 Find the area where contact information is clustered
        Email, phone, and DOB are usually together!
        """
        # Find email position as anchor
        email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        if email_match:
            email_pos = email_match.start()
            # Get 500 chars before and after email
            start = max(0, email_pos - 500)
            end = min(len(text), email_pos + 500)
            return text[start:end]
        
        # Find phone position as anchor
        phone_patterns = [
            r'(?:電話|携帯|TEL|Phone|Mobile)',
            r'\b0[5789]0[-\s]?\d{4}[-\s]?\d{4}\b',
            r'\([0-9]{3}\)',
        ]
        for pattern in phone_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                pos = match.start()
                start = max(0, pos - 500)
                end = min(len(text), pos + 500)
                return text[start:end]
        
        return None

    def _extract_name_english(self, text: str) -> Optional[str]:
        """
        👤 Extract names from ENGLISH resumes!
        Much simpler without Japanese complexity!
        """
        logger.debug("👑 English Name Detective activated!")
        
        text_lines = text.split('\n')
        
        # Words that are DEFINITELY NOT names
        not_names = {
            'resume', 'curriculum', 'vitae', 'cv', 'profile', 'summary',
            'objective', 'experience', 'education', 'skills', 'contact',
            'references', 'work', 'employment', 'professional', 'personal',
            'qualifications', 'achievements', 'certifications', 'projects',
            'engineer', 'developer', 'manager', 'designer', 'analyst',
            'coordinator', 'specialist', 'consultant', 'director', 'senior',
            'junior', 'lead', 'head', 'chief', 'vice', 'assistant'
        }
        
        # Strategy 1: Look for explicit name headers
        name_patterns = [
            r'(?:Name|Full Name|Candidate Name)[\s:]*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})',
            r'^([A-Z][A-Z\s]+)$',  # All caps name on its own line
            r'([A-Z][a-z]+\s+[A-Z]\.\s+[A-Z][a-z]+)',  # John M. Smith
        ]
        
        for pattern in name_patterns:
            matches = re.findall(pattern, text, re.MULTILINE)
            for match in matches:
                candidate = match.strip()
                if self._is_valid_english_name(candidate, not_names):
                    logger.info(f"💎 Found name: {candidate}")
                    return candidate
        
        # Strategy 2: First few lines (names usually at top)
        for i, line in enumerate(text_lines[:10]):
            line = line.strip()
            
            # Skip short, long, or lines with special chars
            if len(line) < 5 or len(line) > 50:
                continue
            
            if any(char in line for char in ['@', 'http', '://', '.com', '|']):
                continue
            
            # Must be title case and 2-4 words
            words = line.split()
            if 2 <= len(words) <= 4:
                if all(word[0].isupper() and word[1:].islower() for word in words if len(word) > 1):
                    if self._is_valid_english_name(line, not_names):
                        logger.info(f"💎 Found name at line {i}: {line}")
                        return line
        
        # Strategy 3: Look near email
        email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        if email_match:
            email_pos = email_match.start()
            context = text[max(0, email_pos-300):email_pos]
            
            name_candidates = re.findall(r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b', context)
            for candidate in reversed(name_candidates):
                if self._is_valid_english_name(candidate, not_names):
                    logger.info(f"💎 Found name near email: {candidate}")
                    return candidate
        
        return None

    def _is_valid_english_name(self, name: str, not_names: set) -> bool:
        """
        ✅ Validate if text is really an English name
        """
        if not name or len(name) < 3 or len(name) > 50:
            return False
        
        name_lower = name.lower()
        
        # Check against blacklist
        if name_lower in not_names:
            return False
        
        if any(word in not_names for word in name_lower.split()):
            return False
        
        # Must be only letters, spaces, hyphens, apostrophes
        if not re.match(r"^[A-Za-z\s\-'\.]+$", name):
            return False
        
        # Should have proper capitalization
        words = name.split()
        for word in words:
            if len(word) > 1:
                # Should start with capital (except Jr, Sr, etc.)
                if word not in ['Jr', 'Sr', 'II', 'III', 'IV'] and not word[0].isupper():
                    return False
        
        # Should not be all caps (unless 2-3 letters)
        if name.isupper() and len(name) > 5:
            return False
        
        return True

    def _is_definitely_a_name(self, text: str, not_names: set) -> bool:
        """
        🔍 ULTRA STRICT validation - is this REALLY a name?
        """
        if not text or len(text) < 2:
            return False
        
        text_lower = text.lower()
        
        # Check against our NOT names list
        if text_lower in not_names:
            return False
        
        # Check each word in multi-word names
        words = text.split()
        for word in words:
            if word.lower() in not_names:
                return False
        
        # Additional validation for English names
        if re.match(r'^[A-Za-z\s\-\.\']+$', text):
            # Should have at least one capital letter
            if not any(c.isupper() for c in text):
                return False
            # Shouldn't be all caps (unless 2-3 chars like "JR")
            if text.isupper() and len(text) > 3:
                return False
            # Each word should start with capital
            for word in words:
                if len(word) > 2 and not word[0].isupper():
                    return False
        
        # Additional validation for Japanese names
        if re.search(r'[一-龯ァ-ヾ]', text):
            # Should not contain numbers or special chars
            if re.search(r'[\d@#$%^&*()_+=\[\]{};:"\\|,.<>/?]', text):
                return False
        
        return True

    def _looks_like_name_part(self, word: str) -> bool:
        """Check if a single word looks like part of a name"""
        # Too short or too long
        if len(word) < 2 or len(word) > 20:
            return False
        # Should start with capital
        if not word[0].isupper():
            return False
        # Shouldn't be all caps (unless short like "JR", "III")
        if word.isupper() and len(word) > 3:
            return False
        # Common name suffixes are OK
        if word.upper() in ['JR', 'SR', 'III', 'II', 'IV']:
            return True
        # Should be mostly letters
        if not re.match(r'^[A-Za-z\'\-\.]+$', word):
            return False
        return True

    def _is_valid_name_strict(self, name: str) -> bool:
        """
        🔍 ULTRA STRICT name validation - NO job titles allowed!
        """
        if not name or len(name) < 2 or len(name) > 50:
            return False
        
        not_names = {
            'profile', 'summary', 'objective', 'experience', 'education',
            'skills', 'references', 'career', 'professional', 'personal',
            'address', 'location', 'portfolio', 'contact', 'information',
            'programmer', 'analyst', 'developer', 'engineer', 'manager',
            'coordinator', 'specialist', 'consultant', 'designer', 'architect',
            'administrator', 'executive', 'director', 'supervisor', 'lead',
            'senior', 'junior', 'intern', 'trainee', 'associate', 'assistant',
            'officer', 'technician', 'expert', 'advisor', 'coach',
            '履歴書', '職務経歴書', '経歴', '学歴', '職歴', 'スキル',
            'プロフィール', 'サマリー', '概要', '自己紹介',
            '日本語能力試験', '試験', 'テスト', '検定'
        }
        
        name_lower = name.lower()
        if name_lower in not_names or any(word in not_names for word in name_lower.split()):
            return False
        
        if name.isupper() and len(name) > 4:
            return False
        
        if any(char in name for char in ['@', 'http', 'www', '/', '\\', '|', '{', '}', '[', ']']):
            return False
        
        if re.match(r'^[A-Za-z\s\-\.\']+$', name):
            if not any(c.isupper() for c in name) or (name.islower() and len(name) > 3):
                return False
        
        return True

    def _clean_name(self, name: str) -> str:
        return re.sub(r'[（(].*?[）)]', '', name).strip()

    def get_text_from_file(self, file_path: str) -> Optional[str]:
        """📄 Enhanced text extraction - ENGLISH ONLY!"""
        file_name = os.path.basename(file_path)
        if file_name.startswith('~$'): 
            return None
        
        text = ""
        try:
            if file_path.lower().endswith('.pdf'):
                # Try pdfplumber first
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                # If no text or very little, try OCR
                if not text or len(text.strip()) < 50:
                    logger.warning(f"📸 Low text from pdfplumber, trying OCR for {file_name}")
                    
                    if self.ocr_available:
                        try:
                            import pytesseract
                            images = pdf2image.convert_from_path(file_path, dpi=300)
                            ocr_text = ""
                            
                            for i, img in enumerate(images):
                                logger.info(f"   🖼️ OCR processing page {i+1}/{len(images)}")
                                page_text = pytesseract.image_to_string(img, lang='eng')
                                ocr_text += page_text + "\n"
                            
                            if len(ocr_text.strip()) > len(text.strip()):
                                text = ocr_text
                                logger.info(f"✅ OCR extracted {len(text)} characters")
                        
                        except Exception as ocr_error:
                            logger.error(f"❌ OCR failed for {file_name}: {ocr_error}")
            
            elif file_path.lower().endswith('.docx'):
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
                
                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += "\n" + cell.text
            
            return self._clean_text(text)
        
        except Exception as e:
            logger.error(f"❌ Extraction failed for {file_name}: {e}")
            return None

    def _extract_data_from_text(self, text: str) -> Tuple[Dict, bool]:
        """
        A helper to extract data from a single block of text.
        This encapsulates our AI-first, Regex-fallback strategy.
        Returns the extracted data and a boolean indicating if AI was used.
        """
        from utils import standardize_phone_number, standardize_date
        final_results = {}
        ai_assisted = False
        
        # Try AI extraction first if enabled
        if self.ai_enabled:
            try:
                ai_results = self.ai_extractor.extract_all_fields(text)
                if ai_results:
                    final_results.update(ai_results)
                    ai_assisted = True
                    logger.info("🤖 AI extraction successful")
            except Exception as e:
                logger.warning(f"⚠️ AI extraction failed: {e}")
        
        # Fallback to regex for any missing fields
        required_fields = ['name', 'email', 'phone', 'date_of_birth', 'skills', 'working_experience', 'location', 'school_university']
        missing_fields = [field for field in required_fields if not final_results.get(field)]
        
        if not self.ai_enabled or missing_fields:
            logger.info(f"⚡ Using regex for fields: {missing_fields if missing_fields else 'all'}")
            regex_results = self._extract_with_mega_regex(text)
            
            # Fill in missing fields from regex results
            for field in missing_fields:
                if not final_results.get(field) and regex_results.get(field):
                    final_results[field] = regex_results.get(field)
        
        # Standardize and format for output
        formatted_data = {
            "Name": final_results.get("name"),
            "Email": final_results.get("email"),
            "Phone": final_results.get("phone"),  # <-- DON'T standardize, keep as-is from AI
            "Date_of_Birth": standardize_date(final_results.get("date_of_birth")),
            "Skills": final_results.get("skills"),
            "Working_Experience": final_results.get("working_experience"),
            "Location": final_results.get("location"),
            "School_University": final_results.get("school_university"),
}
        
        return formatted_data, ai_assisted

    def _extract_name_from_folder(self, folder_path: str) -> Optional[str]:
        """
        💡 Extract name from folder name like "108_Mr ATWAL Prateek"
        """
        folder_name = os.path.basename(folder_path)
        
        # Pattern for folders like "108_Mr ATWAL Prateek" or "109_LEE Edna"
        patterns = [
            r'^\d+_(?:Mr|Ms|Mrs|Dr)?\s*([A-Z][A-Z]+)\s+([A-Za-z]+)',  # 108_Mr ATWAL Prateek
            r'^\d+_([A-Z][A-Z]+)\s+([A-Za-z]+)',  # 109_LEE Edna
            r'^\d+_([A-Za-z]+)\s+([A-Za-z]+)',  # For mixed case
        ]
        
        for pattern in patterns:
            match = re.match(pattern, folder_name)
            if match:
                # Typically Japanese style: LAST FIRST
                last_name = match.group(1)
                first_name = match.group(2) if match.lastindex >= 2 else ""
                
                # Format the name properly
                full_name = f"{first_name} {last_name}".strip()
                if full_name:
                    logger.info(f"📛 Extracted name from folder: {full_name}")
                    return full_name
        
        # Fallback: just remove the ID and clean up
        name_part = re.sub(r'^\d+_', '', folder_name)
        name_part = re.sub(r'(Mr|Ms|Mrs|Dr)\s+', '', name_part)
        name_part = re.sub(r'別名.*', '', name_part).strip()
        
        if name_part and len(name_part) > 2:
            return name_part
        
        return None

    def process_candidate_folder(self, folder_path: str) -> Dict:
        """
        ✨ Processes all resume files - ENGLISH ONLY!
        """
        logger.info(f"🎤 Processing: {os.path.basename(folder_path)}")
        
        # Find resume files
        resume_files = []
        try:
            for filename in os.listdir(folder_path):
                if filename.lower().endswith(('.pdf', '.docx')) and not filename.startswith('~$'):
                    resume_files.append(os.path.join(folder_path, filename))
                    logger.info(f"   ✅ Found: {filename}")
        except Exception as e:
            logger.error(f"❌ Error reading folder: {e}")
            return {}
        
        if not resume_files:
            logger.warning(f"🤷‍♀️ No resume files found")
            return {}
        
        # Initialize result
        result = {
            "ID": None,
            "Name": None,
            "Email": None,
            "Phone": None,
            "Date_of_Birth": None,
            "Skills": None,
            "Working_Experience": None,
            "Location": None,
            "School_University": None,
            "Language": "English",  # ✨ Always English now!
            "Extraction_Status": "Failed",
            "Notes": "",
            "AI_Assisted": False,
            "Filenames_Processed": ", ".join([os.path.basename(f) for f in resume_files])
        }
        
        # Extract ID from folder name
        folder_name = os.path.basename(folder_path)
        id_match = re.match(r'^(\d+)', folder_name)
        if id_match:
            result["ID"] = int(id_match.group(1))
        
        # Combine text from all files
        combined_text = ""
        for file_path in resume_files:
            text = self.get_text_from_file(file_path)
            if text:
                combined_text += text + "\n\n"
        
        # Extract data
        if combined_text and len(combined_text.strip()) > 50:
            extracted_data, ai_used = self._extract_data_from_text(combined_text)
            
            if ai_used:
                result["AI_Assisted"] = True
            
            # Update result
            for field in ["Name", "Email", "Phone", "Date_of_Birth", "Skills", 
                        "Working_Experience", "Location", "School_University"]:
                if extracted_data.get(field):
                    result[field] = extracted_data[field]
        
        # Set status
        self._set_extraction_status(result, combined_text)
        
        logger.info(f"✅ Completed: {result['Extraction_Status']}")
        return result         

    def _extract_name_from_filename(self, file_name: str) -> Optional[str]:
        filename_without_ext = os.path.splitext(file_name)[0]
        filename_patterns = [
            r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})(?:[\s*-_])',
            r'^(?:CV[-_])?([A-Z][a-z]+\s+[A-Z]+)(?:\s+\d|$)',
            r'^([A-Z]+\s+[A-Z][a-z]+)',
            r'^([A-Z][a-z]+\s+[A-Z][a-z]+)'
        ]
        for pattern in filename_patterns:
            match = re.match(pattern, filename_without_ext)
            if match:
                potential_name = match.group(1).strip()
                words = [word.capitalize() if word.isupper() and len(word) > 2 else word for word in potential_name.split()]
                potential_name = ' '.join(words)
                if self._is_valid_name_strict(potential_name):
                    return potential_name
        return None

    def _set_extraction_status(self, result: Dict, text: str):
        if not result["Name"] or not result["ID"]:
            result["Extraction_Status"] = "Failed"
            if not result["Name"]:
                result["Notes"] += "🚨 CRITICAL: No name found! "
            if not result["ID"]:
                result["Notes"] += "🚨 No ID found in folder name!"
        else:
            extracted_count = sum(1 for f in ["ID", "Name", "Email", "Phone", "Date_of_Birth"] if result[f] is not None)
            if extracted_count >= 5:
                result["Extraction_Status"] = "Complete"
            elif extracted_count >= 3:
                result["Extraction_Status"] = "Success"
            else:
                result["Extraction_Status"] = "Partial"

        if not result["Email"] and not result["Phone"]:
            result["Notes"] += " 🚨 NO CONTACT INFO FOUND!"
            self._emergency_contact_extraction(result, text)

    def _emergency_contact_extraction(self, result: Dict, text: str):
        logger.warning(f"🆘 No contact info for {result.get('Filenames_Processed', 'N/A')} - attempting emergency extraction!")
        emergency_patterns = [
            r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
            r'(?:Contact|連絡先|联系).*?(\d[\d\-\s\.]{9,})',
        ]
        for pattern in emergency_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                if '@' in match.group(0) and not result["Email"]:
                    result["Email"] = match.group(1).lower()
                    logger.info(f"🆘 Found emergency email: {result['Email']}")
                elif not result["Phone"]:
                    result["Phone"] = match.group(1)
                    logger.info(f"🆘 Found emergency phone: {result['Phone']}")

    def _normalize_japanese_phone(self, phone: str) -> str:
        phone = phone.translate(str.maketrans('０１２３４５６７８９', '0123456789'))
        phone = re.sub(r'^(\+81|０８１|81)', '0', phone)
        phone = re.sub(r'[^\d\-]', '', phone)
        return phone.replace('・', '-').replace('－', '-').replace(' ', '')

    def _is_valid_japanese_phone(self, phone: str) -> bool:
        digits_only = re.sub(r'\D', '', phone)
        if not 10 <= len(digits_only) <= 11:
            return False
        if re.match(r'^0[5789]0', digits_only):
            return len(digits_only) == 11
        if re.match(r'^0\d{9}$', digits_only):
            return True
        return False

    def _extract_and_validate_phone(self, match: re.Match, pattern_type: str) -> Optional[str]:
        try:
            phone = match.group(0).strip()
            if pattern_type == 'JP':
                return self._normalize_japanese_phone(phone)
            elif pattern_type == 'INTL':
                phone = re.sub(r'[^\d+\-]', '', phone)
                if not phone.startswith('+') and len(re.sub(r'\D', '', phone)) > 10:
                    phone = '+' + phone
                return phone if self._is_valid_international_phone(phone) else None
            elif pattern_type in ['HEADER', 'TABLE', 'BULLET', 'GENERIC']:
                phone = match.group(1) if match.lastindex else phone
                phone = phone.strip()
                if self._looks_like_japanese_phone(phone):
                    return self._normalize_japanese_phone(phone)
                else:
                    phone = re.sub(r'[^\d+\-\s]', '', phone)
                    return phone if len(re.sub(r'\D', '', phone)) >= 10 else None
        except Exception as e:
            logger.debug(f"Phone extraction error: {e}")
        return None
    
    def _looks_like_japanese_phone(self, phone: str) -> bool:
        digits = re.sub(r'\D', '', phone)
        if digits.startswith('81'):
            return True
        if digits.startswith('0') and 10 <= len(digits) <= 11:
            return any(digits.startswith(prefix) for prefix in ['050', '070', '080', '090', '03', '06', '011'])
        return False
    
    def _is_valid_international_phone(self, phone: str) -> bool:
        digits = re.sub(r'\D', '', phone)
        if not 7 <= len(digits) <= 15:
            return False
        if re.match(r'^(\d)\1+$', digits):
            return False
        return True
    
    def _is_sequential(self, digits: str) -> bool:
        if len(digits) < 6: return False
        is_asc = all(int(digits[i]) == int(digits[i-1]) + 1 for i in range(1, min(6, len(digits))))
        is_desc = all(int(digits[i]) == int(digits[i-1]) - 1 for i in range(1, min(6, len(digits))))
        return is_asc or is_desc

    def _extract_from_table_format(self, text: str) -> Dict[str, Optional[str]]:
        logger.debug("📊 Attempting table format extraction!")
        data = {}
        patterns = {
            'name': r'(?:Name|名前|氏名)[\s|│:]\s*([^|│\n]+)',
            'email': r'(?:Email|メール|电子邮件)[\s|│:]\s*([^|│\n]+)',
            'phone': r'(?:Phone|電話|Tel|Mobile)[\s|│:]\s*([^|│\n]+)',
            'date_of_birth': r'(?:DOB|生年月日|Date of Birth)[\s|│:]\s*([^|│\n]+)',
        }
        for key, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                value = match.group(1).strip()
                if key == 'email':
                    data[key] = value.lower()
                else:
                    data[key] = value
        return data

    def check_empty_folders(self, resume_folder: str) -> List[str]:
        """
        🔍 Check for folders without resume files
        """
        empty_folders = []
        
        # Get all candidate folders
        candidate_folders = find_candidate_folders(resume_folder)
        
        for folder_path in candidate_folders:
            # Check if folder has any resume files
            resume_files = find_resume_files(folder_path)
            
            if not resume_files:
                folder_name = os.path.basename(folder_path)
                empty_folders.append(folder_name)
        
        return empty_folders

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'[\u200b\u200c\u200d\ufeff]', '', text)
        text = unicodedata.normalize('NFKC', text)
        return text.replace('｜', '|').replace('—', '-')

    def _extract_with_ai_assistance(self, text: str, regex_results: Dict) -> Dict:
        """
        🤖 Use AI to validate regex results and extract missing fields!
        This is our BACKUP DANCER that helps when regex struggles!
        """
        if not self.ai_enabled:
            return regex_results
        
        logger.info("🤖 AI ASSISTANCE MODE ACTIVATED!")
        
        # Copy results to avoid modifying original
        ai_enhanced_results = regex_results.copy()
        
        # 1️⃣ VALIDATE EXTRACTED NAME
        if regex_results.get("name"):
            # Get context around where we found the name
            name_context = text[:500]  # Usually names are at the top
            is_valid, confidence = self.ai_validator.validate_name(
                regex_results["name"], 
                name_context
            )
            
            if not is_valid or confidence < 0.6:
                logger.warning(f"🚨 AI rejected name '{regex_results['name']}' (confidence: {confidence})")
                # Try to get a better name from AI
                ai_name = self.ai_validator.extract_from_messy_text(text, 'name')
                if ai_name:
                    ai_enhanced_results["name"] = ai_name
                    ai_enhanced_results["ai_extracted_name"] = True
        
        # 2️⃣ FILL IN MISSING FIELDS WITH AI
        missing_fields = []
        if not regex_results.get("name"):
            missing_fields.append("name")
        if not regex_results.get("phone"):
            missing_fields.append("phone")
        if not regex_results.get("date_of_birth"):
            missing_fields.append("dob")
        
        if missing_fields:
            logger.info(f"🤖 AI attempting to extract missing fields: {missing_fields}")
            
            # First, try to fix vertical formatting issues
            if "phone" in missing_fields:
                fixed_text = self._fix_vertical_phone_numbers(text[:2000])
                if fixed_text != text[:2000]:
                    # Re-run phone extraction on fixed text
                    phone = self._extract_phone(fixed_text, regex_results.get("email"))
                    if phone:
                        ai_enhanced_results["phone"] = phone
                        missing_fields.remove("phone")
            
            # Use AI for remaining missing fields
            for field in missing_fields:
                ai_result = self.ai_validator.extract_from_messy_text(text, field)
                if ai_result:
                    if field == "name":
                        ai_enhanced_results["name"] = ai_result
                    elif field == "phone":
                        # Standardize AI-extracted phone
                        ai_enhanced_results["phone"] = self._standardize_phone_format(ai_result, 'AI')
                    elif field == "dob":
                        # Validate date format
                        if re.match(r'\d{4}-\d{2}-\d{2}', ai_result):
                            ai_enhanced_results["date_of_birth"] = ai_result
        
        # 3️⃣ ADD AI CONFIDENCE SCORES
        ai_enhanced_results["ai_assisted"] = True
        ai_enhanced_results["extraction_method"] = "Regex + AI" if any([
            ai_enhanced_results.get("ai_extracted_name"),
            regex_results != ai_enhanced_results
        ]) else "Regex Only"
        
        return ai_enhanced_results

def save_extraction_feedback(self, filename: str, extracted_data: Dict, 
                           corrections: Optional[Dict] = None):
    """
    💾 Save extraction results for future improvement!
    """
    feedback_file = "extraction_feedback.json"
    
    try:
        with open(feedback_file, 'r') as f:
            feedback_data = json.load(f)
    except:
        feedback_data = []
    
    entry = {
        "filename": filename,
        "timestamp": datetime.datetime.now().isoformat(),
        "extracted": extracted_data,
        "corrections": corrections or {},
        "ai_assisted": extracted_data.get("ai_assisted", False)
    }
    
    feedback_data.append(entry)
    
    with open(feedback_file, 'w') as f:
        json.dump(feedback_data, f, indent=2, ensure_ascii=False)
    
    logger.info("💾 Saved extraction feedback for learning!")

def get_user_settings():
    """Gets batch size and max number of resumes to process from the user."""
    while True:
        try:
            batch_size_input = input(f"💅 Batch size? (Enter for {config.DEFAULT_BATCH_SIZE}): ")
            batch_size = int(batch_size_input) if batch_size_input else config.DEFAULT_BATCH_SIZE
            if batch_size <= 0:
                print("😱 Positive numbers only, sweetie!")
                continue

            total_input = input("💅 How many resumes total? (Enter for ALL): ")
            max_to_process = int(total_input) if total_input else None
            if max_to_process is not None and max_to_process <= 0:
                print("😱 Must be positive, darling!")
                continue
            
            return batch_size, max_to_process
        except ValueError:
            print("💔 Numbers only, honey!")

def process_resumes(extractor, folder_list, processed_folders, batch_size):
    """
    Orchestrates the resume processing, batching them for efficiency.
    """
    results = []
    total_processed_resumes = 0
    start_time = time.time()

    performance_monitor = PerformanceMonitor(FeedbackLoopSystem())

    # Process each candidate folder
    total_folders_to_process = len(folder_list)
    pbar = tqdm(total=total_folders_to_process, desc="✨ Processing candidate folders", unit="folder")
    
    for candidate_folder_path in folder_list:
        if candidate_folder_path in processed_folders:
            pbar.update(1)
            continue
        
        # Process this candidate's folder
        result = extractor.process_candidate_folder(candidate_folder_path)
        if result and result.get("ID"):  # Only add if we got a valid result with ID
            results.append(result)
            total_processed_resumes += 1
        
        processed_folders.append(candidate_folder_path)
        save_checkpoint(processed_folders, config.CHECKPOINT_FILE)
        pbar.update(1)
    
    pbar.close()

    end_time = time.time()
    total_duration = end_time - start_time
    logger.info(f"🏁 Processing completed in {total_duration:.2f} seconds.")
    logger.info(f"📊 Total candidates processed: {total_processed_resumes}")
    
    performance_monitor.generate_performance_report()
    return results

def generate_reports(results: List[Dict], empty_folders: List[str]):
    """Generates CSV reports for the extraction results."""
    print("\n" + "═" * 80)
    print("🎉 CURTAIN CALL! The show is complete! 🎉".center(80))
    print("═" * 80)
    
    df = pd.DataFrame(results)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # 🌟 NEW: AI STATISTICS SECTION! 🌟
    # Calculate AI usage stats
    ai_assisted_count = sum(1 for r in results if r.get('AI_Assisted', False))
    ai_enhanced_names = sum(1 for r in results if r.get('AI_Assisted', False) and r.get('Name'))
    ai_enhanced_phones = sum(1 for r in results if r.get('AI_Assisted', False) and r.get('Phone'))
    ai_enhanced_dobs = sum(1 for r in results if r.get('AI_Assisted', False) and r.get('Date_of_Birth'))

    # Main data report
    output_csv = f"{config.OUTPUT_FILENAME_PREFIX}_{timestamp}.csv"
    df.to_csv(output_csv, index=False, encoding='utf-8-sig')
    print(f"\n📄 Main data saved to: {output_csv}")

    # 🗣️ NEW: Language Classification Report 🗣️
    if 'Language' in df.columns:
        print("\n" + "─" * 80)
        print("🗣️ LANGUAGE CLASSIFICATION".center(80))
        print("─" * 80)

        japanese_resumes = df[df['Language'] == 'Japanese']
        english_resumes = df[df['Language'] == 'English']
        unknown_resumes = df[df['Language'] == 'Unknown']

        print(f"\n🇯🇵 Japanese Resumes ({len(japanese_resumes)}):")
        if not japanese_resumes.empty:
            for _, row in japanese_resumes.iterrows():
                print(f"  - ID: {row.get('ID', 'N/A')}, Files: {row.get('Filenames_Processed', 'N/A')}")
        else:
            print("  None found.")

        print(f"\n🇬🇧 English Resumes ({len(english_resumes)}):")
        if not english_resumes.empty:
            for _, row in english_resumes.iterrows():
                print(f"  - ID: {row.get('ID', 'N/A')}, Files: {row.get('Filenames_Processed', 'N/A')}")
        else:
            print("  None found.")

        if not unknown_resumes.empty:
            print(f"\n❓ Unknown Language Resumes ({len(unknown_resumes)}):")
            for _, row in unknown_resumes.iterrows():
                print(f"  - ID: {row.get('ID', 'N/A')}, Files: {row.get('Filenames_Processed', 'N/A')}")

    # Missing fields report
    missing_report = []
    for _, row in df.iterrows():
        missing = [f for f in ['Name', 'Email', 'Phone', 'Date_of_Birth'] if pd.isna(row.get(f))]
        if missing:
            missing_report.append({
                'ID': row.get('ID', 'Unknown'),
                'Filenames_Processed': row['Filenames_Processed'],
                'Missing_Fields': ', '.join(missing)
            })
    
    if missing_report:
        missing_df = pd.DataFrame(missing_report)
        missing_csv = f"missing_fields_report_{timestamp}.csv"
        missing_df.to_csv(missing_csv, index=False, encoding='utf-8-sig')
        print(f"🚨 Missing Fields Report: {missing_csv}")

    # Empty folders report
    if empty_folders:
        empty_df = pd.DataFrame({'Empty_Folders': empty_folders})
        empty_csv = f"empty_folders_report_{timestamp}.csv"
        empty_df.to_csv(empty_csv, index=False, encoding='utf-8-sig')
        print(f"📁 Empty Folders Report: {empty_csv}")

    # Summary statistics
    print("\n📊 EXTRACTION SUMMARY:")
    print(f"✅ Total resumes processed: {len(df)}")
    print(f"🚨 Resumes with missing fields: {len(missing_report)}")
    print(f"📁 Empty candidate folders: {len(empty_folders)}")
    
    # 🌟 AI Assistance Report (OPTIONAL BUT FABULOUS!) 🌟
    if ai_assisted_count > 0:
        ai_report = []
        for _, row in df.iterrows():
            if row.get('AI_Assisted', False):
                ai_report.append({
                    'ID': row.get('ID', 'Unknown'),
                    'Filenames_Processed': row['Filenames_Processed'],
                    'AI_Enhanced_Fields': ', '.join([
                        field for field in ['Name', 'Email', 'Phone', 'Date_of_Birth']
                        if pd.notna(row.get(field))
                    ]),
                    'Extraction_Status': row.get('Extraction_Status', 'Unknown')
                })
        
        if ai_report:
            ai_df = pd.DataFrame(ai_report)
            ai_csv = f"ai_assistance_report_{timestamp}.csv"
            ai_df.to_csv(ai_csv, index=False, encoding='utf-8-sig')
            print(f"\n🤖 AI Assistance Report: {ai_csv}")

    # 🌟 NEW: AI ASSISTANCE STATISTICS! 🌟
    if ai_assisted_count > 0:
        print("\n🤖 AI ASSISTANCE REPORT:")
        print(f"✨ Total AI-Assisted Extractions: {ai_assisted_count}/{len(results)} ({ai_assisted_count/len(results)*100:.1f}%)")
        print(f"   💅 Names enhanced by AI: {ai_enhanced_names}")
        print(f"   📱 Phones extracted by AI: {ai_enhanced_phones}")
        print(f"   🎂 DOBs found by AI: {ai_enhanced_dobs}")
        
        # Which files needed AI help the most?
        ai_files = [r['Filenames_Processed'] for r in results if r.get('AI_Assisted', False)]
        if ai_files and len(ai_files) <= 10:
            print("\n   📁 Files that needed AI help:")
            for filename in ai_files[:5]:  # Show first 5
                print(f"      - {filename}")
            if len(ai_files) > 5:
                print(f"      ... and {len(ai_files) - 5} more")
    else:
        print("\n🤖 AI ASSISTANCE: Not needed - Regex handled everything! 💪")

    print("\n📊 Field Extraction Success Rates:")
    for field in ['ID', 'Name', 'Email', 'Phone', 'Date_of_Birth']:
        if field in df.columns:
            success_rate = (df[field].notna().sum() / len(df)) * 100 if len(df) > 0 else 0
            emoji = "✅" if success_rate > 80 else "⚠️" if success_rate > 50 else "🚨"
            print(f"  {emoji} {field}: {success_rate:.1f}% success")
    
    # 🌟 NEW: Show extraction method breakdown! 🌟
    if 'AI_Assisted' in df.columns:
        print("\n🎯 EXTRACTION METHOD BREAKDOWN:")
        regex_only = len(df[df['AI_Assisted'] == False])
        ai_assisted = len(df[df['AI_Assisted'] == True])
        
        print(f"  ⚡ Regex-only extractions: {regex_only} ({regex_only/len(df)*100:.1f}%)")
        print(f"  🤖 AI-assisted extractions: {ai_assisted} ({ai_assisted/len(df)*100:.1f}%)")
        
        # Performance comparison
        if ai_assisted > 0:
            regex_success = df[df['AI_Assisted'] == False]['Name'].notna().sum() / regex_only * 100 if regex_only > 0 else 0
            ai_success = df[df['AI_Assisted'] == True]['Name'].notna().sum() / ai_assisted * 100
            print(f"\n  📊 Success rate comparison (Name field):")
            print(f"     Regex-only: {regex_success:.1f}%")
            print(f"     AI-assisted: {ai_success:.1f}%")

    # In generate_reports function
    if any(r.get('AI_Used', False) for r in results):
        ai_times = [r.get('processing_time', 0) for r in results if r.get('AI_Used')]
        regex_times = [r.get('processing_time', 0) for r in results if not r.get('AI_Used')]
        
        print("\n🏎️ PERFORMANCE REPORT:")
        print(f"  🌸 AI avg time: {sum(ai_times)/len(ai_times):.2f}s per resume")
        print(f"  ⚡ Regex-only avg time: {sum(regex_times)/len(regex_times):.2f}s per resume")

def find_resume_files(folder_path: str) -> List[str]:
    """
    🔍 Find all resume files (PDF/DOCX) within a folder
    """
    resume_files = []
    
    # Look for PDF and DOCX files
    for item in os.scandir(folder_path):
        if item.is_file() and not item.name.startswith('~$'):
            if item.name.lower().endswith(('.pdf', '.docx')):
                resume_files.append(item.path)
    
    return sorted(resume_files)

def list_candidates_and_files(extractor: "UltimateResumeExtractor"):
    """📜 List all candidates, their resume files, languages, and formats."""
    print("\n" + "📜" * 40)
    print("📜 Candidate File Report 📜".center(80))
    print("📜" * 40 + "\n")

    try:
        all_folders = find_candidate_folders(config.RESUME_FOLDER)
        if not all_folders:
            print(f"😱 No candidate folders found in '{config.RESUME_FOLDER}'!")
            return
    except FileNotFoundError:
        print(f"😱 ERROR: The folder '{config.RESUME_FOLDER}' doesn't exist, darling!")
        return

    candidate_data = {}
    for folder_path in tqdm(all_folders, desc="🕵️‍♀️ Analyzing files"):
        candidate_name = os.path.basename(folder_path)
        
        if candidate_name not in candidate_data:
            candidate_data[candidate_name] = {
                "Candidate": candidate_name,
                "File Details": [],
                "Many Files": "No"
            }

        try:
            files_in_folder = os.listdir(folder_path)
        except FileNotFoundError:
            continue

        if not files_in_folder:
            candidate_data[candidate_name]["File Details"].append("N/A (Folder is empty)")
            continue

        file_count = len(files_in_folder)
        if file_count > 5:
            candidate_data[candidate_name]["Many Files"] = "Yes"

        for filename in files_in_folder:
            file_path = os.path.join(folder_path, filename)
            if os.path.isfile(file_path):
                file_format = os.path.splitext(filename)[1].lower()
                language = "Unknown"
                
                if "japanese" in filename.lower() or any(char in "履歴書職務経歴書" for char in filename):
                    language = "Japanese"
                elif "english" in filename.lower() or "resume" in filename.lower():
                    language = "English"

                candidate_data[candidate_name]["File Details"].append(f"{filename} ({language}, {file_format})")

    report_data = []
    for candidate_name, data in candidate_data.items():
        report_data.append({
            "Candidate": data["Candidate"],
            "File Details": "; ".join(data["File Details"]),
            "Many Files": data["Many Files"]
        })

    if report_data:
        df = pd.DataFrame(report_data)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_csv = f"all_candidates_report_{timestamp}.csv"
        df.to_csv(output_csv, index=False, encoding='utf-8-sig')
        print(f"\n📄 Report saved to: {os.path.abspath(output_csv)}")
        
        # Display the report in the terminal
        print("\n" + df.to_string())
    else:
        print("🤷‍♀️ No candidates found to analyze.")

def main():
    """🎭 The Main Show with MENU MAGIC!"""
    
    from utils import display_menu, save_checkpoint, load_checkpoint, print_batch_table, FeedbackLoopSystem, InteractiveCorrectionSystem, PatternLearningSystem, PerformanceMonitor, select_folders_to_process
    
    folder_list = []
    processed_folders = []
    
    while True:
        display_menu()
        choice = input("\n💅 What's it gonna be, sweetie? (1-7): ").strip()
        
        # 🌟 NEW: Get candidate folders with proper structure! 🌟
        try:
            all_folders = find_candidate_folders(config.RESUME_FOLDER)
            if not all_folders:
                print(f"😱 No candidate folders found in '{config.RESUME_FOLDER}'!")
                print("Make sure your structure is: merlion_resumes/[date_folder]/[candidate_folders]/")
                return
        except FileNotFoundError:
            print(f"😱 ERROR: The folder '{config.RESUME_FOLDER}' doesn't exist, darling!")
            return

        if choice == "1":
            print("\n✨ Starting fresh! Let's process all resumes! ✨")
            print(f"🎯 Found {len(all_folders)} candidate folders to process!")
            processed_folders = []
            folder_list = all_folders
            break
            
        elif choice == "2":
            processed_folders = load_checkpoint(config.CHECKPOINT_FILE)
            if processed_folders:
                print(f"\n👑 Welcome back! Resuming from where you left off with {len(processed_folders)} folders processed.")
            else:
                print("\n🤔 No checkpoint found, darling. Starting fresh!")
            folder_list = [f for f in all_folders if f not in processed_folders]
            break
            
        elif choice == "3":
            print("\n🎯 SELECTIVE PROCESSING")
            
            # Get ALL candidate folders properly
            all_candidate_folders = []
            
            # Scan through date folders
            for date_folder in os.scandir(config.RESUME_FOLDER):
                if date_folder.is_dir() and re.match(r'\d{4}-\d{2}-\d{2}', date_folder.name):
                    print(f"\n📅 Scanning {date_folder.name}...")
                    
                    # Get candidate folders inside this date folder
                    for candidate in os.scandir(date_folder.path):
                        if candidate.is_dir():
                            all_candidate_folders.append(candidate.path)
                            print(f"   ✓ Found: {candidate.name}")
            
            if not all_candidate_folders:
                print("😱 No candidate folders found!")
                continue
            
            print(f"\n💡 Found {len(all_candidate_folders)} total candidate folders")
            print("\nSelect what to process:")
            print("  [1] Process ALL candidates")
            print("  [2] Process specific number of candidates")
            
            sub_choice = input("\nYour choice (1-2): ").strip()
            
            if sub_choice == "1":
                folder_list = all_candidate_folders
            elif sub_choice == "2":
                num = int(input("How many candidates to process? "))
                folder_list = all_candidate_folders[:num]
            else:
                continue
            
            processed_folders = []
            break
            
        elif choice == "4":
            print("\n🔍 DEBUG MODE - Testing ONE folder")
            
            # Find first candidate folder
            test_folder = None
            for date_folder in os.scandir(config.RESUME_FOLDER):
                if date_folder.is_dir() and re.match(r'\d{4}-\d{2}-\d{2}', date_folder.name):
                    for candidate in os.scandir(date_folder.path):
                        if candidate.is_dir():
                            test_folder = candidate.path
                            break
                    if test_folder:
                        break
            
            if test_folder:
                print(f"📁 Testing with: {test_folder}")
                extractor = UltimateResumeExtractor(config.SUZUME_MODEL_NAME)
                result = extractor.process_candidate_folder(test_folder)
                
                print("\n✨ RESULTS:")
                for key, value in result.items():
                    if key != "Filenames_Processed":
                        print(f"  {key}: {value}")
            else:
                print("❌ No test folder found!")
            return
            
        elif choice == "5":
            print("\n📭 Finding candidates with missing resumes...")
            empty_folders = []
            for folder_path in all_folders:
                if not find_resume_files(folder_path):
                    empty_folders.append(os.path.basename(folder_path))
            
            if not empty_folders:
                print("🎉 All candidate folders have resumes! Nothing to report.")
            else:
                print(f"🚨 Found {len(empty_folders)} candidates with missing resumes:")
                for folder_name in empty_folders:
                    print(f"  - {folder_name}")
                
                # Save report
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                report_df = pd.DataFrame({'CandidateFolder_Missing_Resume': empty_folders})
                report_filename = f"missing_resumes_report_{timestamp}.csv"
                report_df.to_csv(report_filename, index=False, encoding='utf-8-sig')
                print(f"\n📄 Report saved to: {report_filename}")
            continue

        elif choice == "7":
            print("\n💋 Until we meet again, darling! Stay fabulous! ✨")
            return

        elif choice == "6":
            print("\n📜 Listing candidates and their resume languages...")
            extractor = UltimateResumeExtractor(config.SUZUME_MODEL_NAME)
            list_candidates_and_files(extractor)
            continue
            
        else:
            print("\n😱 That's not an option, honey! Try again!")

    if not folder_list:
        print("\n🎉 All candidate folders have been processed! Nothing left to do!")
        return

    batch_size, max_to_process = get_user_settings()
    
    extractor = UltimateResumeExtractor(config.SUZUME_MODEL_NAME)
    
    if max_to_process is not None:
        folder_list = folder_list[:max_to_process]
    
    results = process_resumes(extractor, folder_list, processed_folders, batch_size)
    
    if results:
        generate_reports(results, []) # Empty folders check can be integrated elsewhere if needed
    else:
        print("\nNo new data was processed to generate a report.")

if __name__ == "__main__":
    main()